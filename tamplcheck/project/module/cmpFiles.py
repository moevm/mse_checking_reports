# coding=utf-8
"""Библиотека для работы с docx файлами"""
import docx


# анализ формата текста файла
def analysis_doc(filename):
    """ анализ формата текста файла
        :param filename: полный путь до файла + его имя
        :return: список раздичных парметров файла"""
    doc = docx.Document(filename)
    list_with_format = []
    dict_with_format = {}
    for paragraph in doc.paragraphs:
        # Имя стиля
        dict_with_format['Style of paragraph'] = paragraph.style.name
        # Горизонтальное выравнивание
        if paragraph.text != '':
            if paragraph.paragraph_format.alignment is None:
                dict_with_format['Горизонтальное выравнивание'] = doc.styles[paragraph.style.name].paragraph_format.alignment
            else:
                dict_with_format['Горизонтальное выравнивание'] = paragraph.paragraph_format.alignment
        # Отступ слева
        if paragraph.paragraph_format.left_indent is None:
            dict_with_format['Отступ слева'] = doc.styles[paragraph.style.name].paragraph_format.left_indent
        else:
            dict_with_format['Отступ слева'] = paragraph.paragraph_format.left_indent
        # Отступ первой строки абзаца
        if paragraph.paragraph_format.first_line_indent is None:
            dict_with_format['Отступ первой строки абзаца'] = \
                doc.styles[paragraph.style.name].paragraph_format.first_line_indent
        else:
            dict_with_format['Отступ первой строки абзаца'] = paragraph.paragraph_format.first_line_indent
        # Отступ справа
        if paragraph.paragraph_format.right_indent is None:
            dict_with_format['Отступ справа'] = doc.styles[paragraph.style.name].paragraph_format.right_indent
        else:
            dict_with_format['Отступ справа'] = paragraph.paragraph_format.right_indent
        # Табуляция
        # print(paragraph.paragraph_format.tab_stops)
        # Интервалы между абзацами
        if paragraph.paragraph_format.space_before is None:
            dict_with_format['Интервалы между абзацами(перед)'] = doc.styles[
                paragraph.style.name].paragraph_format.space_before
        else:
            dict_with_format['Интервалы между абзацами(перед)'] = paragraph.paragraph_format.space_before
        if paragraph.paragraph_format.space_after is None:
            dict_with_format['Интервалы между абзацами(после)'] = doc.styles[
                paragraph.style.name].paragraph_format.space_after
        else:
            dict_with_format['Интервалы между абзацами(после)'] = paragraph.paragraph_format.space_after
        # Межстрочный интервал
        if paragraph.paragraph_format.line_spacing is None:
            dict_with_format['Межстрочный интервал'] = doc.styles[paragraph.style.name].paragraph_format.line_spacing
        else:
            dict_with_format['Межстрочный интервал'] = paragraph.paragraph_format.line_spacing
        # Поведение при встрече с границей страницы
        # print(paragraph.paragraph_format.keep_together)
        # print(paragraph.paragraph_format.keep_with_next)
        # print(paragraph.paragraph_format.page_break_before)
        # print(paragraph.paragraph_format.widow_control)

        for run in paragraph.runs[:1]:
            # Шрифт
            if run.font.name is None:
                dict_with_format['Шрифт'] = doc.styles[paragraph.style.name].font.name
            else:
                dict_with_format['Шрифт'] = run.font.name
            # Размера шрифта
            if run.font.size is None:
                dict_with_format['Размер текста'] = doc.styles[paragraph.style.name].font.size
            else:
                dict_with_format['Размер текста'] = run.font.size
            # Курсив
            if run.font.italic is None:
                dict_with_format['Курсив'] = doc.styles[paragraph.style.name].font.italic
            else:
                dict_with_format['Курсив'] = run.font.italic
            # Жирный
            if run.font.bold is None:
                dict_with_format['Жирный текст'] = doc.styles[paragraph.style.name].font.bold
            else:
                dict_with_format['Жирный текст'] = run.font.bold
            # Подчеркнутый
            if run.font.underline is None:
                dict_with_format['Подчеркнутый текст'] = doc.styles[paragraph.style.name].font.underline
            else:
                dict_with_format['Подчеркнутый текст'] = run.font.underline
            # Цвет
            if run.font.color.rgb is None:
                dict_with_format['Цвет'] = str(doc.styles[paragraph.style.name].font.color.rgb)
            else:
                dict_with_format['Цвет'] = str(run.font.color.rgb)
        dict_with_format['Text'] = paragraph.text
        list_with_format.append(dict_with_format)
        dict_with_format = {}
    return list_with_format


# Функция проверки на соответиствие шаблону
def comparison_algorithm(template, checked):
    counter_for_template = 0
    counter_for_checked = 0
    result = []
    res = ""
    flag_for_lines = 0
    flag_in_block = 0
    mismatch_list = {}
    true_list = {}
    mismatch_blocks = []
    dict_for_blocks = {}
    # Проверка на структуру
    while counter_for_template < len(template) and counter_for_checked < len(checked):
        if template[counter_for_template].get('Text') == '':
            flag_for_lines = 0
        elif template[counter_for_template].get('Цвет') == '000000' \
                or template[counter_for_template].get('Цвет') == 'None' or template[counter_for_template].get('Цвет') == '00000A':
            flag_for_lines = 0
            k = 0
            while k < len(checked):
                if template[counter_for_template].get('Text') == checked[k].get('Text'):
                    flag_in_block = 1
                    break
                k += 1
            if flag_in_block == 1:
                dict_for_blocks[counter_for_template] = k
                counter_for_checked = k
            else:
                dict_for_blocks[counter_for_template] = -1
                mismatch_blocks.append("Пункт:'{0}', был пропущен, либо в нем была допущена ошибка.".format(template[counter_for_template].get('Text')))
        flag_in_block = 0
        counter_for_template += 1
        counter_for_checked += 1
    if mismatch_blocks:
        mismatch_blocks.append("Исправьте структуру проверямого документа для последущей проверки формата текста.")
        str_b = "\n".join(mismatch_blocks)
        return str_b
    #counter_for_template = 0
    #counter_for_checked = 0
    # Проверка на пустые строки
    #while counter_for_checked < len(template) and counter_for_checked < len(checked):
    #    if template[counter_for_template].get('Text') == '':
    #        if checked[counter_for_checked].get("Text") == '':
    #            counter_for_checked += 1
    #            counter_for_template += 1
    #        else:
    #
    #    elif template[counter_for_template].get('Color') == '000000' or template[counter_for_template].get('Color') == 'None':
    #        mismatch_blocks.append("Была пропущена строка в {0} ".format(counter_for_checked))
    #    else:
    #        while counter_for_template < len(template):
    #            counter_for_checked = dict_for_blocks.get(counter_for_template)
    #            counter_for_template += 1
    flag_in_block = 0
    counter_for_checked = 0
    counter_for_template = 0
    # Проверка на формат
    while counter_for_template < len(template) and counter_for_checked < len(checked):
        if template[counter_for_template].get('Text') == '':
            counter_for_template += 1
        elif template[counter_for_template].get('Цвет') == '000000' \
                or template[counter_for_template].get('Цвет') == 'None' or template[counter_for_template].get('Цвет') == '00000A':
            if dict_for_blocks.get(counter_for_template) != -1:
                counter_for_checked = dict_for_blocks.get(counter_for_template)
                for key in template[counter_for_template].keys():
                    if template[counter_for_template].get(key) != checked[counter_for_checked].get(key):
                        result.append({"В строке {0} шаблона".format(counter_for_template): {key: template[counter_for_template].get(key)}})
                        result.append({"В строке {0} проверяемого документа".format(counter_for_checked): {key: checked[counter_for_checked].get(key)}})
                counter_for_template += 1
                counter_for_checked += 1
            else:
                counter_for_template += 1
        else:
            i = counter_for_template
            while counter_for_template < len(template) and counter_for_checked < len(checked):
                if counter_for_template == len(template) - 1:
                    flag_in_block = len(checked) - 1
                    counter_for_template += 1
                    break
                if dict_for_blocks.get(counter_for_template) is None:
                    counter_for_template += 1
                else:
                    flag_in_block = dict_for_blocks.get(counter_for_template)
                    break
            if template[i].get('Text') != '':
                for j in range(counter_for_checked, flag_in_block):
                    if checked[j].get('Text') != '':
                        for key in checked[j].keys():
                            if key != 'Text' and key != 'Цвет':
                                if template[i].get(key) != checked[j].get(key) and checked[j].get('Text') != '':
                                    mismatch_list[key] = checked[j].get(key)
                                    true_list[key] = template[i].get(key)
                        if mismatch_list:
                            result.append({"В строке {0} шаблона".format(counter_for_template - 1): true_list,
                                           "В строке {0} проверяемого документа".format(j): mismatch_list})
                        mismatch_list = {}
                        true_list = {}
    if not result:
        return "Проверяемый файл соответствует шаблону."
    else:
        for s in result:
            res = res + str(s) + "\n"
        return res


# Считывание файлов
#template_doc = docx.Document("test.docx")
#checked_doc = docx.Document("test1.docx")
#
# словарь, в котором будем хранить информацию о формате текста документа
#template_format = analysis_doc(template_doc)
#checked_format = analysis_doc(checked_doc)
# вывод полей
#for i in range(0, len(checked_format), 2):
#   for j in checked_format[i]:
#       print(j, ":", checked_format[i].get(j))
# вывод отличий файла
#print(comparison_algorithm(template_format, checked_format))